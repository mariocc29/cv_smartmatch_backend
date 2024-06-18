from django.http import HttpResponse
from django.template.loader import get_template
from rest_framework.generics import GenericAPIView
from rest_framework.response import Response
from rest_framework import status
from xhtml2pdf import pisa
from django.utils.translation import activate
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from common.models.job_offer_model import JobOfferModel
from web.services.resume.handler import ResumeHandler

class ResumePDFView(GenericAPIView):

  def get(self, request, *args, **kwargs):
    try:
      job_offer = JobOfferModel.resume(kwargs['job_offer_id'])
    except StopIteration:
      return Response({'error': 'Job offer not found'}, status=status.HTTP_404_NOT_FOUND)
    
    lang = request.GET.get('lang', 'en')
    activate(lang)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'filename="{job_offer['personal_info']['fullname']}.pdf"'

    template = get_template('resume.html')
    html = template.render(ResumeHandler.data(job_offer, lang))
    pisa_status = pisa.CreatePDF(html, dest=response, encoding='utf-8')

    if pisa_status.err:
      return HttpResponse({'error': 'Error generating PDF'}, status=500)
    
    return response

class ResumeDOCView(GenericAPIView):
  def get(self, request, *args, **kwargs):
    try:
      job_offer = JobOfferModel.resume(kwargs['job_offer_id'])
    except StopIteration:
      return Response({'error': 'Job offer not found'}, status=status.HTTP_404_NOT_FOUND)
    
    lang = request.GET.get('lang', 'en')
    activate(lang)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'filename="{job_offer['personal_info']['fullname']}.docx"'

    template = get_template('resume.html')
    html = template.render(ResumeHandler.data(job_offer, lang))
    
    soup = BeautifulSoup(html, 'html.parser')
    document = Document()
    
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.save(response)
    return response